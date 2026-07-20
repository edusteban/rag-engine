from pathlib  import Path 
from langchain_core.documents import Document

from pypdf import PdfReader
from docx import Document as DocxDocument

def load_txt(path: Path) -> list[Document]:
    with open (path, "r", encoding="utf-8") as f:
        content = f.read()
        return [Document(
            page_content=content, 
        )]
    
def load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(path)

    documents = []

    for page in reader.pages:
        content = page.extract_text()

        if not content:
            continue

        documents.append(
            Document(
                page_content=content
            )
        )

    return documents


def load_docx(path: Path) -> list[Document]:
    doc = DocxDocument(path)

    content = "\n".join(
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    )

    return [
        Document(
            page_content=content
        )
    ]

def load_md(path: Path) -> list[Document]:
    return load_txt(path)
